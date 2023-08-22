import streamlit as st
import pandas as pd
import asyncio
import time
import io
import base64
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import database as db
import columns_categories_config as ccconfig

@st.cache_data
def fetch_all_herbs_cache():
    data = db.fetch_all_herbs()
    return data


def add_new_herb(herb_id, brand, herb_name, cost_price, selling_price, stock):
    async def add_new_herb_async(herb_id, brand, herb_name, cost_price, selling_price, stock):
        db.insert_herb(herb_id, brand, herb_name, cost_price, selling_price, stock)
        st.cache_data.clear()

    coro = add_new_herb_async(herb_id, brand, herb_name, cost_price, selling_price, stock)
    with st.spinner("正在加入中藥數據..."):
        asyncio.run(coro)
    st.success(ccconfig.SUCCESS_MSG)
    time.sleep(1)
    st.experimental_rerun()


def update_inventory(old_df, edited_df):
    different_rows = (old_df != edited_df).any(axis=1)
    herb_id = old_df[different_rows]['key'].values[0]
    for col in edited_df.columns:
        if edited_df[col][different_rows].values[0] != old_df[col][different_rows].values[0]:
            col_changed = col
            new_value = edited_df.loc[different_rows, col].values[0]
            if not isinstance(new_value, str):
                new_value = float(new_value)
            break

    async def update_herb_async(herb_id, col_changed, new_value):
        db.update_herb(herb_id, {f"{col_changed}": new_value})
        st.cache_data.clear()

    coro = update_herb_async(herb_id, col_changed, new_value)
    with st.spinner("正在更新中藥數據..."):
        asyncio.run(coro)
    st.success(ccconfig.SUCCESS_MSG)
    time.sleep(1)
    st.experimental_rerun()


def remove_herb(herb_id):
    async def remove_herb_async(herb_id):
        db.delete_herb(herb_id)
        st.cache_data.clear()
    coro = remove_herb_async(herb_id)
    with st.spinner("正在移除中藥數據..."):
        asyncio.run(coro)
    st.success(ccconfig.SUCCESS_MSG)
    time.sleep(1)
    st.experimental_rerun()



def inventory():
    BRANDS = ccconfig.HERB_BRANDS # ["Sam Gau", "Hoi Tin", "Others"]
    COLUMN_ORDER = ccconfig.INVENTORY_COLUMN_ORDER # ["key", "brand", "herb_name", "cost_price", "selling_price", "inventory"]
    inventory_data = fetch_all_herbs_cache()
    df_inventory = pd.DataFrame(inventory_data, columns=COLUMN_ORDER) # initialize dataframe with the expected column order
    herb_id_list = df_inventory['key'].tolist() # list of all herb_id

    # Display input fields for adding new inventory entries
    st.header("新增中藥")
    herb_id = st.text_input("中藥編號")
    brand = st.selectbox("品牌", BRANDS)
    herb_name = st.text_input("中藥名稱")
    cost_price = st.number_input("來貨價", step=0.1)
    selling_price = st.number_input("零售價", step=0.1) # Can be 0 or None
    stock = st.number_input("數量", step=1)

    if st.button("新增中藥"):
        if any(field == "" for field in [herb_id, brand, herb_name, cost_price, stock]):
            st.warning(ccconfig.WARNING_MSG_FILL_ALL)
        elif herb_id in herb_id_list:
            st.warning("中藥編號已存在，請填寫另一個編號。")
        else:
            add_new_herb(herb_id, brand, herb_name, cost_price, selling_price, stock)


        
    # Display the current inventory table (using st.data_editor)
    st.divider()
    st.header("現時存貨")

    col1, col2 = st.columns([2,1])
    with col1:
        st.subheader("編輯中藥數據：")
        selected_brand = st.multiselect(
        "按品牌篩選",
        options=BRANDS,
        default=BRANDS,
        )
        df_inventory = df_inventory[
            (df_inventory['brand'].isin(selected_brand))
        ]
        edited_df_inventory = st.data_editor(
            df_inventory, 
            use_container_width=True,
            num_rows="fixed",
            hide_index=True,
            column_config = {
                "key": st.column_config.Column("中藥編號", disabled=True, help=ccconfig.INFO_MSG_NOT_EDITABLE),
                "brand": st.column_config.TextColumn("品牌", disabled=True, help=ccconfig.INFO_MSG_NOT_EDITABLE),
                "herb_name": st.column_config.TextColumn("中藥名稱"),
                "cost_price": st.column_config.NumberColumn("來貨價", min_value=0, format="$%d"),
                "selling_price": st.column_config.NumberColumn("零售價", min_value=0, format="$%d"),
                "inventory": st.column_config.NumberColumn("數量", min_value=0, step=1),
            },
        )
        if edited_df_inventory is not None and not edited_df_inventory.equals(df_inventory):
            update_inventory(df_inventory, edited_df_inventory)
    
    with col2:
        st.subheader("移除中藥：")
        herb_id_list = df_inventory["key"].tolist()
        chosen_herb_id = st.text_input("輸入中藥編號移除：")
        with st.expander("確認移除中藥", expanded=False):
            delete_button = st.button("確認", type="primary")

        if delete_button:
            if chosen_herb_id:
                if chosen_herb_id in herb_id_list:
                    remove_herb(chosen_herb_id)
                    st.success(f"已移除中藥編號：{chosen_herb_id}，現在現在刷新頁面...")
                    time.sleep(1)
                    st.experimental_rerun()
                else:
                    st.warning("此中藥編號不存在。")
            else:
                st.warning("請選擇中藥編號。")
                

    st.divider()
    for brand in BRANDS:
        brand_data = df_inventory[df_inventory["brand"] == brand]
        brand_data = brand_data[COLUMN_ORDER]
        brand_data = brand_data.drop(columns=["brand"])
        num_items = brand_data.shape[0]  # Count number of items for the brand
        st.subheader(f"{brand} ({num_items}項)")
        st.dataframe(
            brand_data, 
            hide_index=True,
            use_container_width=True,
            column_config={
                "key": st.column_config.Column("中藥編號", disabled=True),
                "herb_name": st.column_config.TextColumn("中藥名稱"),
                "cost_price": st.column_config.NumberColumn("來貨價", format="$%d"),
                "selling_price": st.column_config.NumberColumn("零售價", min_value=0, format="$%d"),
                "inventory": st.column_config.NumberColumn("數量"),
            }
        )

    # Export excel and word files
    st.divider()
    st.subheader("下載存貨文件")
    arrange_by_mapping = {
        "key": "中藥編號",
        "herb_name": "中藥名稱"
    }
    arrange_by = st.selectbox("排列方式", ["key", "herb_name"], format_func=lambda option: arrange_by_mapping[option])

    # Export Excel button
    excel_dataframes = {
        f"{brand}存貨": df_inventory[df_inventory["brand"] == brand][COLUMN_ORDER]
        .drop(columns=["brand"])
        .sort_values(by=arrange_by)
        .reset_index(drop=True)
        .rename(columns={"key": "中藥編號", "herb_name": "中藥名稱", "cost_price": "來貨價", "selling_price": "零售價", "inventory": "數量"})
        for brand in BRANDS
    }
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
        for sheet_name, data in excel_dataframes.items():
            data.to_excel(writer, sheet_name=sheet_name, index=False)
    b64_excel = base64.b64encode(excel_buffer.getvalue()).decode()
    href_excel = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64_excel}" download="inventory.xlsx" class="button">下載Excel報告</a>'
    st.markdown(href_excel, unsafe_allow_html=True)

    # Export Word button
    doc = docx.Document()
    for brand, data in excel_dataframes.items():
        num_items = data.shape[0]  # Count number of items for the brand
        doc.add_heading(f"{brand} ({num_items}項)", level=1)  # Add num_items to the title
        # Add table with gridlines
        table = doc.add_table(data.shape[0] + 1, data.shape[1])
        table.style = "Table Grid"
        for j, col_name in enumerate(data.columns):
            table.cell(0, j).text = col_name
            for i in range(data.shape[0]):
                table.cell(i + 1, j).text = str(data.iloc[i, j])
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    b64_word = base64.b64encode(doc_buffer.read()).decode()
    href_word = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_word}" download="inventory.docx" class="button">下載Word報告</a>'
    st.markdown(href_word, unsafe_allow_html=True)


if __name__ == "__main__":
    inventory()
