import streamlit as st
import pandas as pd
import utils
import matplotlib.pyplot as plt
from matplotlib import font_manager
import io
import base64
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import database as db
import columns_categories_config as ccconfig

@st.cache_data
def fetch_all_incomes_cached():
    data = db.fetch_all_incomes()
    return data

@st.cache_data
def fetch_all_expenses_cached():
    data = db.fetch_all_expenses()
    return data

def convert_to_monthly_summary(df, month_options):
    if df.empty:
        return pd.DataFrame(columns=["month", "amount"])  # Return an empty DataFrame with the required columns
    summary_df = df.groupby('month', as_index=False)['amount'].sum()
    summary_df['month'] = pd.Categorical(summary_df['month'], categories=month_options, ordered=True)
    summary_df = summary_df.sort_values('month')
    return summary_df[['month', 'amount']] 


@st.cache_data
def get_income_expense_by_month(df_income_by_month, df_expense_by_month):
    df_income_expense_by_month = pd.merge(df_income_by_month, df_expense_by_month, on='month', how='outer')
    df_income_expense_by_month = df_income_expense_by_month.rename(columns={
        'amount_x': 'Income',
        'amount_y': 'Expense'
    })
    months_list = df_income_expense_by_month['month'].tolist()
    return df_income_expense_by_month, months_list


def create_word_report(dataframes):
    doc = Document()
    for title, df in dataframes.items():
        doc.add_heading(title, level=1)
        # Add table with gridlines
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = "Table Grid"
        # Set column headers
        for j, col_name in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = col_name
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Fill data into the table
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cell = table.cell(i + 1, j)
                cell.text = str(df.iat[i, j])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc
    

def statistics():
    st.header("統計")
    
    income_data = fetch_all_incomes_cached()
    expense_data = fetch_all_expenses_cached()

    col1, col2 = st.columns(2)
    with col1:
        INCOME_CATEGORIES = ccconfig.INCOME_CATEGORIES # ["Consultation", "Herb Sale", "Class", "Others"]
        INCOME_COLUMN_ORDER = ccconfig.INCOME_COLUMN_ORDER # ["key", "category", "item", "customer", "amount"]
        df_income = pd.DataFrame(income_data, columns=INCOME_COLUMN_ORDER)
        df_income['month'] = df_income['key'].apply(utils.format_month) # add a new column "month", by reading the date from "key"
        month_options = df_income['month'].unique() # list of months, e.g. ['2023 May' '2023 Jun' '2023 Jul' '2023 Aug']
        # Display for checking only, will hide it
        # st.dataframe(df_income,
        #                 hide_index=True, 
        #                 use_container_width=True,
        #                 column_config={
        #                 "key": st.column_config.Column("Income ID", disabled=True, help="Info: Not editable"),
        #                 "category": st.column_config.TextColumn("Category"),
        #                 "item": st.column_config.TextColumn("Item"),
        #                 "customer": st.column_config.TextColumn("Customer"),
        #                 "amount": st.column_config.NumberColumn("Amount", format="$%d"),
        #                 }
        #             )
    if len(income_data) > 0:
        df_income_by_category = df_income.pivot_table(index='month', columns='category', values='amount', aggfunc='sum', fill_value=0)
        df_income_by_category.reset_index(inplace=True)
        df_income_by_category = df_income_by_category.set_index('month').reindex(month_options).reset_index()
        df_income_by_category = df_income_by_category[['month'] + INCOME_CATEGORIES]
    else:
        df_income_by_category = pd.DataFrame(columns=['month'] + INCOME_CATEGORIES)

    
    st.subheader("收入類別")
    st.dataframe(df_income_by_category, 
                    hide_index=True,
                    use_container_width=True,
                    column_config={"month": st.column_config.TextColumn("月份")}
                )


    with col2:
        EXPENSE_CATEGORIES = ccconfig.EXPENSE_CATEGORIES # ["Rent", "Salaries", "Utilities", "Advertising", "Travel", "Others"]
        EXPENSE_COLUMN_ORDER = ccconfig.EXPENSE_COLUMN_ORDER # ["key", "category", "item", "amount"]
        df_expense = pd.DataFrame(expense_data, columns=EXPENSE_COLUMN_ORDER)
        df_expense['month'] = df_expense['key'].apply(utils.format_month) # add a new column "month", by reading the date from "key"
        month_options = df_expense['month'].unique() # list of months, e.g. ['2023 May' '2023 Jun' '2023 Jul' '2023 Aug']
        # Display for checking only, will hide it
        # st.dataframe(df_expense,
        #                 hide_index=True, 
        #                 use_container_width=True,
        #                 column_config={
        #                 "key": st.column_config.Column("Expense ID", disabled=True, help="Info: Not editable"),
        #                 "category": st.column_config.TextColumn("Category"),
        #                 "item": st.column_config.TextColumn("Item"),
        #                 "amount": st.column_config.NumberColumn("Amount", format="$%d"),
        #                 }
        #             )
    if len(expense_data) > 0:
        df_expense_by_category = df_expense.pivot_table(index='month', columns='category', values='amount', aggfunc='sum', fill_value=0)
        df_expense_by_category.reset_index(inplace=True)
        df_expense_by_category = df_expense_by_category.set_index('month').reindex(month_options).reset_index()
        df_expense_by_category = df_expense_by_category[['month'] + EXPENSE_CATEGORIES]
    else:
        df_expense_by_category = pd.DataFrame(columns=['month'] + EXPENSE_CATEGORIES)


    st.subheader("支出類別")
    st.dataframe(df_expense_by_category,
                    hide_index=True,
                    use_container_width=True,
                    column_config={"month": st.column_config.TextColumn("月份")},
                )


    if len(income_data) > 0 and len(expense_data) > 0:
        st.subheader("總收入支出總結")
        df_income_by_month = convert_to_monthly_summary(df_income, month_options)
        df_income_by_month = df_income_by_month.reset_index(drop=True)
        df_expense_by_month = convert_to_monthly_summary(df_expense, month_options)
        df_expense_by_month = df_expense_by_month.reset_index(drop=True)

        df_income_expense_by_month, months_list = get_income_expense_by_month(df_income_by_month, df_expense_by_month)
        df_income_expense_by_month["Net Income"] = df_income_expense_by_month["Income"] - df_income_expense_by_month["Expense"]
        st.dataframe(df_income_expense_by_month,
                        hide_index=True,
                        use_container_width=True,
                        column_config={
                        "month": st.column_config.TextColumn("月份"),
                        "Income": st.column_config.NumberColumn("總收入", format="$%d"),
                        "Expense": st.column_config.NumberColumn("總支出", format="$%d"),
                        "Net Income": st.column_config.NumberColumn("淨收入", format="$%d"),
                        }
                    )
        


        st.divider()
        st.subheader("收入支出趨勢")
        start_month, end_month = st.select_slider(
            "選擇時限：",
            options=months_list,
            value=(months_list[0], months_list[-1]) 
        )

        # Filter Line chart of "Income by category" by selected months
        df_income_by_category['month'] = pd.to_datetime(df_income_by_category['month'], format='%Y %b')
        df_income_by_category = df_income_by_category[
            (df_income_by_category['month'] >= pd.to_datetime(start_month, format='%Y %b')) & 
            (df_income_by_category['month'] <= pd.to_datetime(end_month, format='%Y %b'))
        ]
        df_income_by_category['month'] = df_income_by_category['month'].dt.strftime('%Y %b')

        # Filter # Line chart of "Expense by category" by selected months
        df_expense_by_category['month'] = pd.to_datetime(df_expense_by_category['month'], format='%Y %b')
        df_expense_by_category = df_expense_by_category[
            (df_expense_by_category['month'] >= pd.to_datetime(start_month, format='%Y %b')) & 
            (df_expense_by_category['month'] <= pd.to_datetime(end_month, format='%Y %b'))
        ]
        df_expense_by_category['month'] = df_expense_by_category['month'].dt.strftime('%Y %b')

        # Filter Line chart of "Total Income, Expense and Net Income" by selected months
        df_income_expense_by_month = df_income_expense_by_month[
            (df_income_expense_by_month['month'] >= start_month) & (df_income_expense_by_month['month'] <= end_month)
        ]

        # Set Chinese font
        fontP = font_manager.FontProperties(fname="./fonts/SimHei.ttf")
        fontP.set_family('SimHei') 
        fontP.set_size(14)

        # Display Line chart of "Income by category"
        plt.figure(figsize=(10, 4))
        plt.title('收入類別', fontproperties=fontP)
        for category in INCOME_CATEGORIES:
            plt.plot(df_income_by_category['month'], df_income_by_category[category], marker='^', label=category)
        plt.xlabel('月份', fontproperties=fontP)
        plt.ylabel('金額', fontproperties=fontP)
        plt.legend(loc=0, prop=fontP)
        plt.grid(True)
        st.pyplot(plt)

        # Display Line chart of "Expense by category"
        plt.figure(figsize=(10, 4))
        plt.title('支出類別', fontproperties=fontP)
        for category in EXPENSE_CATEGORIES:
            plt.plot(df_expense_by_category['month'], df_expense_by_category[category], marker='v', label=category)
        plt.xlabel('月份', fontproperties=fontP)
        plt.ylabel('金額', fontproperties=fontP)
        plt.legend(loc=2, prop=fontP) # '2' means 'upper left'
        plt.grid(True)
        st.pyplot(plt)

        # Display Line chart of "Total Income, Expense and Net Income"
        plt.figure(figsize=(10, 4))
        plt.title('總收入、總支出和淨收入', fontproperties=fontP)
        plt.plot(df_income_expense_by_month['month'], df_income_expense_by_month['Income'], marker='^', label='總收入')
        plt.plot(df_income_expense_by_month['month'], df_income_expense_by_month['Expense'], marker='v', label='總支出')
        plt.plot(df_income_expense_by_month['month'], df_income_expense_by_month['Net Income'], marker='o', label='淨收入')
        plt.xlabel('月份', fontproperties=fontP)
        plt.ylabel('金額', fontproperties=fontP)
        plt.legend(loc=0, prop=fontP)
        plt.grid(True)
        st.pyplot(plt)


        # Export excel and word files
        st.divider()
        st.subheader("下載統計文件")
        # Export excel file
        excel_buffer = io.BytesIO()

        # Modify names of columns
        df1_2_new_column_names = {'month': '月份'}
        df_income_by_category.rename(columns=df1_2_new_column_names, inplace=True)
        df_expense_by_category.rename(columns=df1_2_new_column_names, inplace=True)

        df3_new_column_names = {'month': '月份', 'Income': '總收入', 'Expense': '總支出', 'Net Income': '淨收入'}
        df_income_expense_by_month.rename(columns=df3_new_column_names, inplace=True)

        with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
            df_income_by_category.to_excel(writer, sheet_name="收入類別", index=False)
            df_expense_by_category.to_excel(writer, sheet_name="支出類別", index=False)
            df_income_expense_by_month.to_excel(writer, sheet_name="總收入支出總結", index=False)

        b64 = base64.b64encode(excel_buffer.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="statistics.xlsx" class="button">下載Excel報告</a>'
        st.markdown(href, unsafe_allow_html=True)

        # Export word file
        dataframes = {
            "收入類別": df_income_by_category,
            "支出類別": df_expense_by_category,
            "總收入支出總結": df_income_expense_by_month
        }
        doc = create_word_report(dataframes)
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        b64 = base64.b64encode(doc_buffer.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="statistics.docx" class="button">下載Word報告</a>'
        st.markdown(href, unsafe_allow_html=True)

    else:
        st.subheader("總收入支出總結")
        st.warning("請注意！您的收入或支出數據並不存在。 如果想查看總收入、總支出及相關圖表，請先引入相關數據。")


if __name__ == "__main__":
    statistics()